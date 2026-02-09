"""
Resume parsing service for extracting structured data from PDF and DOCX files.

This module provides functionality to parse resumes in various formats and
extract structured information using AI.
"""

import logging
from typing import Dict, Any, Optional
from io import BytesIO
import PyPDF2
from docx import Document
from anthropic import Anthropic
import os
import json

logger = logging.getLogger(__name__)


class ResumeParserService:
    """Service for parsing resume files and extracting structured data."""

    def __init__(self):
        """Initialize the resume parser service."""
        self.anthropic_client = Anthropic(api_key=os.getenv("ANTHROPIC_API_KEY"))

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file.

        Args:
            file_content: PDF file content as bytes

        Returns:
            str: Extracted text from the PDF

        Raises:
            Exception: If PDF extraction fails
        """
        try:
            pdf_file = BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            logger.info(f"Extracted {len(text)} characters from PDF")
            return text.strip()

        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {str(e)}")
            raise Exception(f"Failed to parse PDF file: {str(e)}")

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_content: DOCX file content as bytes

        Returns:
            str: Extracted text from the DOCX

        Raises:
            Exception: If DOCX extraction fails
        """
        try:
            docx_file = BytesIO(file_content)
            doc = Document(docx_file)

            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + "\n"

            logger.info(f"Extracted {len(text)} characters from DOCX")
            return text.strip()

        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {str(e)}")
            raise Exception(f"Failed to parse DOCX file: {str(e)}")

    def parse_resume_with_ai(self, resume_text: str) -> Dict[str, Any]:
        """
        Parse resume text using Claude AI to extract structured data.

        Args:
            resume_text: Raw text extracted from resume

        Returns:
            Dict containing structured resume data

        Raises:
            Exception: If AI parsing fails
        """
        try:
            prompt = f"""You are a resume parsing expert. Extract structured information from the following resume text.

CRITICAL: Preserve the EXACT wording, keywords, and technical terms as written. Do NOT paraphrase, summarize, or reword anything.

Resume Text:
{resume_text}

Extract and return ONLY a valid JSON object with the following structure. Do not include any other text or explanation:

{{
    "full_name": "extracted full name",
    "email": "extracted email",
    "phone": "extracted phone number",
    "location": "extracted location/city",
    "summary": "extracted professional summary or objective",
    "experience": [
        {{
            "company": "company name",
            "position": "job title",
            "start_date": "start date",
            "end_date": "end date or Present",
            "bullets": ["achievement or responsibility 1", "achievement or responsibility 2", "achievement or responsibility 3"]
        }}
    ],
    "education": [
        {{
            "institution": "school/university name",
            "degree": "degree type and field",
            "start_date": "start date",
            "end_date": "end date",
            "gpa": "GPA if mentioned"
        }}
    ],
    "skills": ["skill1", "skill2", "skill3"],
    "projects": [
        {{
            "name": "project name",
            "description": "project description",
            "technologies": ["tech1", "tech2"]
        }}
    ],
    "certifications": ["certification1", "certification2"],
    "languages": ["language1", "language2"]
}}

CRITICAL INSTRUCTIONS:
- Copy text VERBATIM - do not paraphrase or summarize
- Preserve ALL keywords, metrics, and technical terms exactly as written
- Extract every bullet point completely - do not shorten or combine them
- Keep all numbers, percentages, and quantifications exactly as stated
- Maintain the exact phrasing of achievements (e.g., "Led", "Developed", "Implemented")
- If any field is not found, use null or empty array []
- Ensure all dates are in readable format
- Extract ALL skills mentioned, including technical tools, frameworks, and methodologies
- For experience bullets: extract EVERY achievement separately, preserving exact wording
- Return ONLY the JSON object, no additional text"""

            response = self.anthropic_client.messages.create(
                model="claude-sonnet-4-5-20250929",  # Claude 4.5 Sonnet
                max_tokens=8192,  # Increased for better parsing
                messages=[{
                    "role": "user",
                    "content": prompt
                }]
            )

            # Extract JSON from response
            response_text = response.content[0].text.strip()

            # Remove markdown code blocks if present (more robust)
            if "```json" in response_text:
                # Extract content between ```json and ```
                start = response_text.find("```json") + 7
                end = response_text.find("```", start)
                response_text = response_text[start:end].strip()
            elif "```" in response_text:
                # Extract content between ``` and ```
                start = response_text.find("```") + 3
                end = response_text.find("```", start)
                response_text = response_text[start:end].strip()

            response_text = response_text.strip()

            # Parse JSON
            parsed_data = json.loads(response_text)

            logger.info("Successfully parsed resume with AI")
            return parsed_data

        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse AI response as JSON: {str(e)}")
            logger.error(f"Response text: {response_text}")
            raise Exception("Failed to parse AI response. Please try again.")

        except Exception as e:
            logger.error(f"Failed to parse resume with AI: {str(e)}")
            raise Exception(f"Failed to parse resume: {str(e)}")

    def parse_resume_file(
        self,
        file_content: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """
        Parse a resume file (PDF or DOCX) and extract structured data.

        Args:
            file_content: File content as bytes
            filename: Original filename

        Returns:
            Dict containing structured resume data

        Raises:
            ValueError: If file format is not supported
            Exception: If parsing fails
        """
        # Determine file type
        file_extension = filename.lower().split(".")[-1]

        # Extract text based on file type
        if file_extension == "pdf":
            resume_text = self.extract_text_from_pdf(file_content)
        elif file_extension in ["docx", "doc"]:
            resume_text = self.extract_text_from_docx(file_content)
        else:
            raise ValueError(
                f"Unsupported file format: {file_extension}. "
                "Please upload PDF or DOCX files only."
            )

        # Validate extracted text
        if not resume_text or len(resume_text) < 50:
            raise ValueError(
                "Could not extract meaningful text from the file. "
                "Please ensure the file is not empty or corrupted."
            )

        # Parse with AI
        parsed_data = self.parse_resume_with_ai(resume_text)

        return parsed_data

    def convert_to_resume_format(self, parsed_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Convert parsed data to our resume database format.

        Args:
            parsed_data: Parsed data from AI

        Returns:
            Dict in our resume content format
        """
        # Build resume content matching our schema (using camelCase for consistency)
        content = {
            "personalInfo": {
                "name": parsed_data.get("full_name", ""),
                "email": parsed_data.get("email", ""),
                "phone": parsed_data.get("phone", ""),
                "location": parsed_data.get("location", ""),
                "linkedin": "",
                "github": ""
            },
            "summary": parsed_data.get("summary") or "",
            "experience": [],
            "education": [],
            "skills": parsed_data.get("skills") or [],
            "projects": [],
            "certifications": parsed_data.get("certifications") or [],
            "languages": parsed_data.get("languages") or []
        }

        # Convert experience (handle None values)
        for exp in parsed_data.get("experience") or []:
            # Combine dates into duration string for PDF
            start = exp.get("start_date", "")
            end = exp.get("end_date", "")
            duration = f"{start} - {end}" if start and end else ""

            # Get bullets array, or convert description to single bullet if present
            bullets = exp.get("bullets") or []
            if not bullets and exp.get("description"):
                bullets = [exp.get("description")]

            content["experience"].append({
                "company": exp.get("company", ""),
                "title": exp.get("position", ""),  # Changed from "position" to "title" to match frontend
                "duration": duration,
                "bullets": bullets
            })

        # Convert education (handle None values)
        for edu in parsed_data.get("education") or []:
            # Combine dates into duration string for PDF
            start = edu.get("start_date", "")
            end = edu.get("end_date", "")
            duration = f"{start} - {end}" if start and end else ""

            content["education"].append({
                "institution": edu.get("institution", ""),
                "degree": edu.get("degree", ""),
                "duration": duration,
                "gpa": edu.get("gpa", "")
            })

        # Convert projects (handle None values)
        for proj in parsed_data.get("projects") or []:
            content["projects"].append({
                "name": proj.get("name", ""),
                "description": proj.get("description", ""),
                "technologies": proj.get("technologies") or [],
                "link": ""
            })

        return content


# Singleton instance
resume_parser_service = ResumeParserService()
